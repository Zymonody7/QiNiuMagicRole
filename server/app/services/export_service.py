"""
导出服务 - 处理对话导出功能
"""

import asyncio
import io
import tempfile
import os
from typing import List, Dict, Any, Optional
from datetime import datetime
import requests
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from pydub import AudioSegment
from pydub.effects import normalize
import aiofiles

from app.core.config import settings
from app.services.tts_service import TTSService
from app.services.static_asset_service import static_asset_service
from app.services.qiniu_podcast_tts_service import qiniu_podcast_tts_service


class ExportService:
    """导出服务类"""
    
    def __init__(self, db):
        self.db = db
        self.tts_service = TTSService()
        self.storage_service = static_asset_service
    
    def _safe_string(self, text):
        """安全处理字符串，避免编码问题"""
        if text is None:
            return ""
        if isinstance(text, str):
            # 清理可能的编码问题
            return text.encode('utf-8', errors='ignore').decode('utf-8')
        return str(text)
    
    async def generate_text_export(
        self, 
        messages: List[Dict], 
        character: Any, 
        format_type: str = "word"
    ) -> bytes:
        """生成文本导出文件"""
        if format_type == "word":
            return await self._generate_word_document(messages, character)
        elif format_type == "pdf":
            return await self._generate_pdf_document(messages, character)
        else:
            raise ValueError(f"不支持的格式: {format_type}")
    
    async def _generate_word_document(self, messages: List[Dict], character: Any) -> bytes:
        """生成Word文档"""
        doc = Document()
        
        # 设置文档标题
        title = doc.add_heading(f'与{character.name}的对话记录', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加角色信息
        doc.add_heading('角色信息', level=1)
        doc.add_paragraph(f'角色名称: {self._safe_string(character.name)}')
        doc.add_paragraph(f'角色描述: {self._safe_string(character.description)}')
        doc.add_paragraph(f'性格特点: {self._safe_string(character.personality)}')
        doc.add_paragraph(f'背景故事: {self._safe_string(character.background)}')
        
        # 添加对话记录
        doc.add_heading('对话记录', level=1)
        doc.add_paragraph(f'导出时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'总消息数: {len(messages)}')
        doc.add_paragraph('')
        
        # 添加消息内容
        for i, message in enumerate(messages, 1):
            # 消息头部
            speaker = "用户" if message.get('is_user', False) else character.name
            timestamp = message.get('created_at', '')
            if timestamp:
                try:
                    # 如果timestamp已经是datetime对象，直接格式化
                    if isinstance(timestamp, datetime):
                        timestamp = timestamp.strftime('%Y-%m-%d %H:%M:%S')
                    else:
                        # 如果是字符串，尝试解析
                        dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                        timestamp = dt.strftime('%Y-%m-%d %H:%M:%S')
                except:
                    # 如果解析失败，尝试截取字符串
                    if isinstance(timestamp, str):
                        timestamp = timestamp[:19] if len(timestamp) > 19 else timestamp
                    else:
                        timestamp = str(timestamp)[:19]
            
            doc.add_heading(f'{i}. {speaker} ({timestamp})', level=2)
            
            # 消息内容
            content = message.get('content', '')
            # 确保内容为字符串并处理编码
            if isinstance(content, str):
                # 清理可能的编码问题
                content = content.encode('utf-8', errors='ignore').decode('utf-8')
            else:
                content = str(content)
            doc.add_paragraph(content)
            doc.add_paragraph('')  # 空行分隔
        
        # 保存到内存
        doc_buffer = io.BytesIO()
        try:
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            return doc_buffer.getvalue()
        except UnicodeEncodeError as e:
            # 如果出现编码错误，尝试重新处理文档内容
            print(f"编码错误，尝试修复: {e}")
            # 重新创建文档，确保所有内容都是安全的
            return await self._generate_word_document_safe(messages, character)
    
    async def _generate_word_document_safe(self, messages: List[Dict], character: Any) -> bytes:
        """安全生成Word文档，处理编码问题"""
        doc = Document()
        
        # 设置文档标题
        title = doc.add_heading(f'与{self._safe_string(character.name)}的对话记录', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加角色信息
        doc.add_heading('角色信息', level=1)
        doc.add_paragraph(f'角色名称: {self._safe_string(character.name)}')
        doc.add_paragraph(f'角色描述: {self._safe_string(character.description)}')
        doc.add_paragraph(f'性格特点: {self._safe_string(character.personality)}')
        doc.add_paragraph(f'背景故事: {self._safe_string(character.background)}')
        
        # 添加对话记录
        doc.add_heading('对话记录', level=1)
        doc.add_paragraph(f'导出时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'总消息数: {len(messages)}')
        doc.add_paragraph('')
        
        # 添加消息内容
        for i, message in enumerate(messages, 1):
            # 消息头部
            speaker = "用户" if message.get('is_user', False) else self._safe_string(character.name)
            timestamp = message.get('created_at', '')
            if timestamp:
                try:
                    if isinstance(timestamp, datetime):
                        timestamp = timestamp.strftime('%Y-%m-%d %H:%M:%S')
                    else:
                        dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                        timestamp = dt.strftime('%Y-%m-%d %H:%M:%S')
                except:
                    if isinstance(timestamp, str):
                        timestamp = timestamp[:19] if len(timestamp) > 19 else timestamp
                    else:
                        timestamp = str(timestamp)[:19]
            
            doc.add_heading(f'{i}. {speaker} ({timestamp})', level=2)
            
            # 消息内容
            content = self._safe_string(message.get('content', ''))
            doc.add_paragraph(content)
            doc.add_paragraph('')
        
        # 保存到内存
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer.getvalue()
    
    async def _generate_pdf_document(self, messages: List[Dict], character: Any) -> bytes:
        """生成PDF文档"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # 注册中文字体
        chinese_font_name = 'Helvetica'  # 默认字体
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # 常见的中文字体路径
            chinese_fonts = [
                'C:/Windows/Fonts/simsun.ttc',  # 宋体
                'C:/Windows/Fonts/simhei.ttf',  # 黑体
                'C:/Windows/Fonts/msyh.ttc',    # 微软雅黑
                '/System/Library/Fonts/PingFang.ttc',  # macOS
                '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',  # Linux
            ]
            
            for font_path in chinese_fonts:
                try:
                    if os.path.exists(font_path):
                        pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                        chinese_font_name = 'ChineseFont'
                        print(f"成功注册中文字体: {font_path}")
                        break
                except Exception as e:
                    print(f"注册字体失败 {font_path}: {e}")
                    continue
                    
        except Exception as e:
            print(f"字体注册失败: {e}")
        
        # 创建自定义样式
        chinese_heading1 = ParagraphStyle(
            'ChineseHeading1',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # 居中
            fontName=chinese_font_name
        )
        
        chinese_heading2 = ParagraphStyle(
            'ChineseHeading2',
            parent=styles['Heading2'],
            fontName=chinese_font_name
        )
        
        chinese_normal = ParagraphStyle(
            'ChineseNormal',
            parent=styles['Normal'],
            fontName=chinese_font_name
        )
        
        chinese_heading3 = ParagraphStyle(
            'ChineseHeading3',
            parent=styles['Heading3'],
            fontName=chinese_font_name
        )
        
        # 添加标题
        story.append(Paragraph(f'与{self._safe_string(character.name)}的对话记录', chinese_heading1))
        story.append(Spacer(1, 20))
        
        # 角色信息
        story.append(Paragraph('角色信息', chinese_heading2))
        story.append(Paragraph(f'<b>角色名称:</b> {self._safe_string(character.name)}', chinese_normal))
        story.append(Paragraph(f'<b>角色描述:</b> {self._safe_string(character.description)}', chinese_normal))
        story.append(Paragraph(f'<b>性格特点:</b> {self._safe_string(character.personality)}', chinese_normal))
        story.append(Paragraph(f'<b>背景故事:</b> {self._safe_string(character.background)}', chinese_normal))
        story.append(Spacer(1, 20))
        
        # 对话记录
        story.append(Paragraph('对话记录', chinese_heading2))
        story.append(Paragraph(f'导出时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', chinese_normal))
        story.append(Paragraph(f'总消息数: {len(messages)}', chinese_normal))
        story.append(Spacer(1, 20))
        
        # 添加消息内容
        for i, message in enumerate(messages, 1):
            speaker = "用户" if message.get('is_user', False) else self._safe_string(character.name)
            timestamp = message.get('created_at', '')
            if timestamp:
                try:
                    if isinstance(timestamp, datetime):
                        timestamp = timestamp.strftime('%Y-%m-%d %H:%M:%S')
                    else:
                        dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                        timestamp = dt.strftime('%Y-%m-%d %H:%M:%S')
                except:
                    if isinstance(timestamp, str):
                        timestamp = timestamp[:19] if len(timestamp) > 19 else timestamp
                    else:
                        timestamp = str(timestamp)[:19]
            
            # 消息标题
            story.append(Paragraph(f'{i}. {speaker} ({timestamp})', chinese_heading3))
            
            # 消息内容
            content = self._safe_string(message.get('content', ''))
            content = content.replace('\n', '<br/>')
            story.append(Paragraph(content, chinese_normal))
            story.append(Spacer(1, 12))
        
        try:
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
        except UnicodeEncodeError as e:
            print(f"PDF编码错误，尝试修复: {e}")
            return await self._generate_pdf_document_safe(messages, character)
    
    async def _generate_pdf_document_safe(self, messages: List[Dict], character: Any) -> bytes:
        """安全生成PDF文档，处理编码问题"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # 标题样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # 居中
        )
        
        # 添加标题
        story.append(Paragraph(f'与{self._safe_string(character.name)}的对话记录', title_style))
        story.append(Spacer(1, 20))
        
        # 角色信息
        story.append(Paragraph('角色信息', styles['Heading2']))
        story.append(Paragraph(f'<b>角色名称:</b> {self._safe_string(character.name)}', styles['Normal']))
        story.append(Paragraph(f'<b>角色描述:</b> {self._safe_string(character.description)}', styles['Normal']))
        story.append(Paragraph(f'<b>性格特点:</b> {self._safe_string(character.personality)}', styles['Normal']))
        story.append(Paragraph(f'<b>背景故事:</b> {self._safe_string(character.background)}', styles['Normal']))
        story.append(Spacer(1, 20))
        
        # 对话记录
        story.append(Paragraph('对话记录', styles['Heading2']))
        story.append(Paragraph(f'导出时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', styles['Normal']))
        story.append(Paragraph(f'总消息数: {len(messages)}', styles['Normal']))
        story.append(Spacer(1, 20))
        
        # 添加消息内容
        for i, message in enumerate(messages, 1):
            speaker = "用户" if message.get('is_user', False) else self._safe_string(character.name)
            timestamp = message.get('created_at', '')
            if timestamp:
                try:
                    if isinstance(timestamp, datetime):
                        timestamp = timestamp.strftime('%Y-%m-%d %H:%M:%S')
                    else:
                        dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                        timestamp = dt.strftime('%Y-%m-%d %H:%M:%S')
                except:
                    if isinstance(timestamp, str):
                        timestamp = timestamp[:19] if len(timestamp) > 19 else timestamp
                    else:
                        timestamp = str(timestamp)[:19]
            
            # 消息标题
            story.append(Paragraph(f'{i}. {speaker} ({timestamp})', styles['Heading3']))
            
            # 消息内容
            content = self._safe_string(message.get('content', ''))
            content = content.replace('\n', '<br/>')
            story.append(Paragraph(content, styles['Normal']))
            story.append(Spacer(1, 12))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    async def generate_podcast_audio(
        self, 
        messages: List[Dict], 
        character: Any,
        background_music: Optional[str] = None
    ) -> bytes:
        """生成播客音频"""
        try:
            print(f"开始生成播客，共 {len(messages)} 条消息")
            
            # 创建临时目录
            with tempfile.TemporaryDirectory() as temp_dir:
                audio_segments = []
                
                # 添加播客开头
                intro_audio = await self._generate_podcast_intro(character)
                if intro_audio:
                    audio_segments.append(intro_audio)
                
                # 处理每条消息 - 按时间顺序
                for i, message in enumerate(messages):
                    print(f"处理消息 {i+1}/{len(messages)} - {'用户' if message.get('is_user', False) else 'AI'}")
                    content = message.get('content', '')
                    is_user = message.get('is_user', False)
                    existing_audio_url = message.get('audio_url')
                    
                    # 不再添加说话者标识，直接播放内容
                    
                    if existing_audio_url and not is_user:
                        # 使用现有的AI音频
                        try:
                            audio_segment = await self._download_audio(existing_audio_url)
                            if audio_segment:
                                # 标准化音频
                                audio_segment = self._normalize_audio(audio_segment)
                                audio_segments.append(audio_segment)
                        except Exception as e:
                            print(f"下载现有音频失败: {e}")
                            # 生成新的音频
                            audio_segment = await self._generate_tts_audio(content, character, is_user)
                            if audio_segment:
                                # 对用户音频使用增强的标准化
                                if is_user:
                                    audio_segment = self._normalize_audio_with_gain(audio_segment, target_db=-18.0)
                                else:
                                    audio_segment = self._normalize_audio(audio_segment)
                                audio_segments.append(audio_segment)
                    else:
                        # 生成TTS音频
                        audio_segment = await self._generate_tts_audio(content, character, is_user)
                        if audio_segment:
                            # 对用户音频使用增强的标准化
                            if is_user:
                                audio_segment = self._normalize_audio_with_gain(audio_segment, target_db=-18.0)
                            else:
                                audio_segment = self._normalize_audio(audio_segment)
                            audio_segments.append(audio_segment)
                    
                    # 添加消息间隔
                    if i < len(messages) - 1:  # 不是最后一条消息
                        pause_audio = self._generate_pause()
                        if pause_audio:
                            audio_segments.append(pause_audio)
                
                if not audio_segments:
                    raise Exception("没有生成任何音频片段")
                
                print(f"开始拼接 {len(audio_segments)} 个音频片段")
                
                # 拼接所有音频，使用音量平衡
                final_audio = self._concatenate_audios_with_volume_balance(audio_segments)
                
                # 添加播客结尾
                outro_audio = await self._generate_podcast_outro(character)
                if outro_audio:
                    final_audio = final_audio + outro_audio
                
                # 添加背景音乐（可选）
                final_audio = await self._add_background_music(final_audio, background_music)
                
                # 最终音频处理
                final_audio = self._finalize_audio(final_audio)
                
                # 导出为MP3
                output_path = os.path.join(temp_dir, "podcast.mp3")
                final_audio.export(output_path, format="mp3", bitrate="128k")
                
                print(f"播客生成完成，文件大小: {os.path.getsize(output_path)} 字节")
                
                # 读取文件内容
                async with aiofiles.open(output_path, 'rb') as f:
                    return await f.read()
                    
        except Exception as e:
            print(f"生成播客音频失败: {e}")
            import traceback
            traceback.print_exc()
            raise Exception(f"播客生成失败: {str(e)}")
    
    async def _download_audio(self, audio_url: str) -> Optional[AudioSegment]:
        """下载音频文件"""
        try:
            print(f"尝试下载AI音频: {audio_url}")
            response = requests.get(audio_url, timeout=30)
            if response.status_code == 200:
                print(f"音频下载成功，大小: {len(response.content)} 字节")
                # 创建临时文件
                with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as temp_file:
                    temp_file.write(response.content)
                    temp_file.flush()
                    temp_file_path = temp_file.name
                
                try:
                    # 加载音频
                    audio = AudioSegment.from_file(temp_file_path)
                    print(f"AI音频加载成功，时长: {len(audio)}ms")
                    
                    # 检查音频是否有声音
                    if len(audio) > 0:
                        # 获取音频的RMS（均方根）值，用于判断是否有声音
                        rms = audio.rms
                        print(f"音频RMS值: {rms}")
                        if rms < 100:  # 如果RMS值很小，可能是静音
                            print("警告：音频可能为静音或音量过低")
                        else:
                            print("音频音量正常")
                    else:
                        print("警告：音频时长为0")
                    
                    return audio
                finally:
                    # 确保删除临时文件
                    try:
                        os.unlink(temp_file_path)
                    except:
                        pass
            else:
                print(f"音频下载失败，状态码: {response.status_code}")
                return None
        except Exception as e:
            print(f"下载音频失败: {e}")
            return None
    
    async def _generate_tts_audio(self, text: str, character: Any, is_user: bool) -> Optional[AudioSegment]:
        """生成TTS音频"""
        try:
            # 优先使用七牛云播客TTS服务
            if qiniu_podcast_tts_service.is_enabled():
                try:
                    # 根据用户类型选择不同的音色
                    if is_user:
                        # 用户使用七牛云男声
                        voice_type = "qiniu_zh_male_whxkxg"  # 温和学科小哥
                    else:
                        # AI角色使用女声（作为备用）
                        voice_type = "qiniu_zh_female_wwxkjx"
                    
                    print(f"使用七牛云播客TTS生成音频: {text[:30]}...")
                    audio = await qiniu_podcast_tts_service.generate_podcast_voice(
                        text=text,
                        voice_type=voice_type,
                        speed_ratio=1.0,
                        encoding="mp3"
                    )
                    
                    if audio:
                        print(f"七牛云播客TTS生成成功，时长: {len(audio)}ms")
                        return audio
                    else:
                        print("七牛云播客TTS生成失败，降级到模拟音频")
                        return self._generate_mock_audio(text, is_user)
                        
                except Exception as e:
                    print(f"七牛云播客TTS生成失败: {e}")
                    # 降级到模拟音频
                    return self._generate_mock_audio(text, is_user)
            else:
                print("七牛云播客TTS服务未启用，使用模拟音频")
                return self._generate_mock_audio(text, is_user)
            
        except Exception as e:
            print(f"生成TTS音频失败: {e}")
            return self._generate_mock_audio(text, is_user)
    
    def _concatenate_audios(self, audio_segments: List[AudioSegment]) -> AudioSegment:
        """拼接音频片段"""
        if not audio_segments:
            raise Exception("没有音频片段可拼接")
        
        # 添加短暂静音分隔
        silence = AudioSegment.silent(duration=500)  # 0.5秒静音
        
        result = audio_segments[0]
        for segment in audio_segments[1:]:
            result += silence + segment
        
        return result
    
    def _concatenate_audios_with_volume_balance(self, audio_segments: List[AudioSegment]) -> AudioSegment:
        """拼接音频片段并平衡音量"""
        if not audio_segments:
            raise Exception("没有音频片段可拼接")
        
        # 计算所有音频片段的平均音量
        total_db = 0
        valid_segments = []
        
        for segment in audio_segments:
            if segment.dBFS != float('-inf'):  # 跳过静音片段
                total_db += segment.dBFS
                valid_segments.append(segment)
        
        if not valid_segments:
            return audio_segments[0] if audio_segments else AudioSegment.silent(duration=1000)
        
        # 计算目标音量（平均音量）
        target_db = total_db / len(valid_segments)
        print(f"音频拼接 - 目标音量: {target_db:.1f}dB")
        
        # 添加短暂静音分隔
        silence = AudioSegment.silent(duration=500)  # 0.5秒静音
        
        result = None
        for i, segment in enumerate(audio_segments):
            # 调整每个片段的音量到目标音量
            if segment.dBFS != float('-inf'):
                volume_diff = target_db - segment.dBFS
                # 限制音量调整范围，避免过度调整
                volume_diff = max(-10, min(10, volume_diff))
                if abs(volume_diff) > 1:  # 只调整差异较大的音频
                    segment = segment + volume_diff
                    print(f"音频片段 {i+1} 音量调整: {volume_diff:.1f}dB")
            
            if result is None:
                result = segment
            else:
                result += silence + segment
        
        return result
    
    async def _add_background_music(self, audio: AudioSegment, background_music: Optional[str] = None) -> AudioSegment:
        """添加背景音乐（可选功能）"""
        try:
            if not background_music:
                print("未选择背景音乐，跳过背景音乐添加")
                return audio
            
            print(f"添加背景音乐: {background_music}")
            
            # 根据背景音乐类型选择不同的音乐
            music_segment = await self._get_background_music_segment(background_music)
            if music_segment:
                # 调整背景音乐长度以匹配主音频
                if len(music_segment) < len(audio):
                    # 如果背景音乐太短，循环播放
                    loops_needed = (len(audio) // len(music_segment)) + 1
                    music_segment = music_segment * loops_needed
                
                # 截取到主音频长度
                music_segment = music_segment[:len(audio)]
                
                # 降低背景音乐音量（-10dB，比之前大一些）
                music_segment = music_segment - 10
                
                # 混合音频
                mixed_audio = audio.overlay(music_segment)
                print(f"背景音乐混合完成，原时长: {len(audio)}ms，混合后: {len(mixed_audio)}ms")
                return mixed_audio
            else:
                print("背景音乐加载失败，使用原音频")
                return audio
                
        except Exception as e:
            print(f"添加背景音乐失败: {e}")
            return audio
    
    async def _generate_podcast_intro(self, character: Any) -> Optional[AudioSegment]:
        """生成播客开头"""
        try:
            intro_text = "欢迎收听对话播客。"
            return await self._generate_tts_audio(intro_text, character, False)
        except Exception as e:
            print(f"生成播客开头失败: {e}")
            return None
    
    async def _generate_podcast_outro(self, character: Any) -> Optional[AudioSegment]:
        """生成播客结尾"""
        try:
            outro_text = "感谢收听对话播客，再见！"
            return await self._generate_tts_audio(outro_text, character, False)
        except Exception as e:
            print(f"生成播客结尾失败: {e}")
            return None
    
    
    def _generate_pause(self) -> AudioSegment:
        """生成暂停音效"""
        try:
            # 减少到0.5秒的静音，让对话更流畅
            return AudioSegment.silent(duration=500)
        except Exception as e:
            print(f"生成暂停失败: {e}")
            return AudioSegment.silent(duration=300)
    
    def _normalize_audio(self, audio: AudioSegment) -> AudioSegment:
        """标准化音频"""
        try:
            # 标准化音量
            normalized = normalize(audio)
            # 确保音频长度至少500ms
            if len(normalized) < 500:
                normalized = normalized + AudioSegment.silent(duration=500 - len(normalized))
            return normalized
        except Exception as e:
            print(f"音频标准化失败: {e}")
            return audio
    
    def _normalize_audio_with_gain(self, audio: AudioSegment, target_db: float = -20.0) -> AudioSegment:
        """标准化音频并调整到目标音量"""
        try:
            # 计算当前音频的RMS音量
            current_db = audio.dBFS
            
            # 如果音频太安静，增加音量
            if current_db < target_db - 10:  # 如果比目标音量低10dB以上
                gain_db = target_db - current_db
                # 限制最大增益为20dB，避免爆音
                gain_db = min(gain_db, 20.0)
                audio = audio + gain_db
                print(f"音频音量调整: {current_db:.1f}dB -> {audio.dBFS:.1f}dB (增益: {gain_db:.1f}dB)")
            
            # 标准化音量
            normalized = normalize(audio)
            # 确保音频长度至少500ms
            if len(normalized) < 500:
                normalized = normalized + AudioSegment.silent(duration=500 - len(normalized))
            return normalized
        except Exception as e:
            print(f"音频标准化失败: {e}")
            return audio
    
    def _finalize_audio(self, audio: AudioSegment) -> AudioSegment:
        """最终音频处理"""
        try:
            # 添加淡入淡出效果
            fade_duration = min(1000, len(audio) // 10)  # 淡入淡出时间
            
            # 淡入
            audio = audio.fade_in(fade_duration)
            # 淡出
            audio = audio.fade_out(fade_duration)
            
            # 标准化音量
            audio = normalize(audio)
            
            return audio
        except Exception as e:
            print(f"最终音频处理失败: {e}")
            return audio
    
    async def _get_background_music_segment(self, music_type: str) -> Optional[AudioSegment]:
        """获取背景音乐片段"""
        try:
            # 根据音乐类型生成不同的背景音乐
            if music_type == "soft":
                # 柔和背景音乐 - 生成柔和的音调
                return self._generate_soft_music()
            elif music_type == "ambient":
                # 环境音乐 - 生成环境音效
                return self._generate_ambient_music()
            elif music_type == "classical":
                # 古典音乐 - 生成古典风格音调
                return self._generate_classical_music()
            elif music_type == "jazz":
                # 爵士音乐 - 生成爵士风格音调
                return self._generate_jazz_music()
            else:
                # 默认柔和音乐
                return self._generate_soft_music()
                
        except Exception as e:
            print(f"获取背景音乐失败: {e}")
            return None
    
    def _generate_soft_music(self) -> AudioSegment:
        """生成柔和背景音乐"""
        try:
            import numpy as np
            
            # 生成柔和的音调序列
            sample_rate = 44100
            duration = 30000  # 30秒
            t = np.linspace(0, duration / 1000, int(sample_rate * duration / 1000), False)
            
            # 生成多个柔和音调
            wave1 = np.sin(2 * np.pi * 220 * t) * 0.1  # A3
            wave2 = np.sin(2 * np.pi * 330 * t) * 0.08  # E4
            wave3 = np.sin(2 * np.pi * 440 * t) * 0.06  # A4
            
            # 组合音调
            combined_wave = wave1 + wave2 + wave3
            
            # 添加淡入淡出
            fade_samples = int(sample_rate * 2)  # 2秒淡入淡出
            combined_wave[:fade_samples] *= np.linspace(0, 1, fade_samples)
            combined_wave[-fade_samples:] *= np.linspace(1, 0, fade_samples)
            
            # 转换为AudioSegment
            audio_data = (combined_wave * 16383).astype(np.int16)
            return AudioSegment(
                audio_data.tobytes(),
                frame_rate=sample_rate,
                sample_width=2,
                channels=1
            )
        except Exception as e:
            print(f"生成柔和音乐失败: {e}")
            return AudioSegment.silent(duration=30000)
    
    def _generate_ambient_music(self) -> AudioSegment:
        """生成环境音乐"""
        try:
            import numpy as np
            
            sample_rate = 44100
            duration = 30000  # 30秒
            t = np.linspace(0, duration / 1000, int(sample_rate * duration / 1000), False)
            
            # 生成环境音效 - 低频嗡嗡声
            wave = np.sin(2 * np.pi * 60 * t) * 0.05  # 低频
            wave += np.sin(2 * np.pi * 120 * t) * 0.03  # 中低频
            
            # 添加随机噪声模拟环境音
            noise = np.random.normal(0, 0.01, len(t))
            combined_wave = wave + noise
            
            # 添加淡入淡出
            fade_samples = int(sample_rate * 3)  # 3秒淡入淡出
            combined_wave[:fade_samples] *= np.linspace(0, 1, fade_samples)
            combined_wave[-fade_samples:] *= np.linspace(1, 0, fade_samples)
            
            audio_data = (combined_wave * 16383).astype(np.int16)
            return AudioSegment(
                audio_data.tobytes(),
                frame_rate=sample_rate,
                sample_width=2,
                channels=1
            )
        except Exception as e:
            print(f"生成环境音乐失败: {e}")
            return AudioSegment.silent(duration=30000)
    
    def _generate_classical_music(self) -> AudioSegment:
        """生成古典音乐"""
        try:
            import numpy as np
            
            sample_rate = 44100
            duration = 30000  # 30秒
            t = np.linspace(0, duration / 1000, int(sample_rate * duration / 1000), False)
            
            # 生成古典和声
            wave1 = np.sin(2 * np.pi * 261.63 * t) * 0.08  # C4
            wave2 = np.sin(2 * np.pi * 329.63 * t) * 0.06  # E4
            wave3 = np.sin(2 * np.pi * 392.00 * t) * 0.05  # G4
            
            combined_wave = wave1 + wave2 + wave3
            
            # 添加淡入淡出
            fade_samples = int(sample_rate * 2)
            combined_wave[:fade_samples] *= np.linspace(0, 1, fade_samples)
            combined_wave[-fade_samples:] *= np.linspace(1, 0, fade_samples)
            
            audio_data = (combined_wave * 16383).astype(np.int16)
            return AudioSegment(
                audio_data.tobytes(),
                frame_rate=sample_rate,
                sample_width=2,
                channels=1
            )
        except Exception as e:
            print(f"生成古典音乐失败: {e}")
            return AudioSegment.silent(duration=30000)
    
    def _generate_jazz_music(self) -> AudioSegment:
        """生成爵士音乐"""
        try:
            import numpy as np
            
            sample_rate = 44100
            duration = 30000  # 30秒
            t = np.linspace(0, duration / 1000, int(sample_rate * duration / 1000), False)
            
            # 生成爵士和声
            wave1 = np.sin(2 * np.pi * 220 * t) * 0.07  # A3
            wave2 = np.sin(2 * np.pi * 277.18 * t) * 0.05  # C#4
            wave3 = np.sin(2 * np.pi * 329.63 * t) * 0.04  # E4
            
            combined_wave = wave1 + wave2 + wave3
            
            # 添加淡入淡出
            fade_samples = int(sample_rate * 2)
            combined_wave[:fade_samples] *= np.linspace(0, 1, fade_samples)
            combined_wave[-fade_samples:] *= np.linspace(1, 0, fade_samples)
            
            audio_data = (combined_wave * 16383).astype(np.int16)
            return AudioSegment(
                audio_data.tobytes(),
                frame_rate=sample_rate,
                sample_width=2,
                channels=1
            )
        except Exception as e:
            print(f"生成爵士音乐失败: {e}")
            return AudioSegment.silent(duration=30000)
    
    def _generate_mock_audio(self, text: str, is_user: bool) -> AudioSegment:
        """生成模拟音频（当TTS不可用时）"""
        try:
            # 根据文本长度计算音频时长（大约每10个字符1秒）
            duration = max(1000, len(text) * 100)  # 最少1秒，每字符100ms
            
            # 生成不同频率的音调来区分用户和AI
            if is_user:
                # 用户使用较高频率
                frequency = 440  # A4音符
            else:
                # AI使用较低频率
                frequency = 220  # A3音符
            
            # 生成正弦波音频
            import numpy as np
            sample_rate = 44100
            t = np.linspace(0, duration / 1000, int(sample_rate * duration / 1000), False)
            wave = np.sin(2 * np.pi * frequency * t) * 0.1  # 低音量
            
            # 添加淡入淡出
            fade_samples = int(sample_rate * 0.1)  # 0.1秒淡入淡出
            wave[:fade_samples] *= np.linspace(0, 1, fade_samples)
            wave[-fade_samples:] *= np.linspace(1, 0, fade_samples)
            
            # 转换为AudioSegment
            audio_data = (wave * 32767).astype(np.int16)
            audio = AudioSegment(
                audio_data.tobytes(),
                frame_rate=sample_rate,
                sample_width=2,
                channels=1
            )
            
            print(f"生成模拟音频: {text[:20]}... (时长: {duration}ms)")
            return audio
            
        except Exception as e:
            print(f"生成模拟音频失败: {e}")
            # 返回静音
            return AudioSegment.silent(duration=2000)
